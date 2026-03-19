import io
import os
import subprocess
import tempfile
import logging
import requests
import base64
import mimetypes
import asyncio
import re
import hashlib
import concurrent.futures

from typing import Tuple, List, Optional, Dict, Any, TYPE_CHECKING

# === LIGHTWEIGHT PARSERS (all in requirements.txt) ===
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
from charset_normalizer import detect as detect_encoding
import tiktoken
from bs4 import BeautifulSoup

# === UNSTRUCTURED (hi_res strategy for complex PDFs) ===
from unstructured.partition.auto import partition
from unstructured.partition.pdf import partition_pdf
from unstructured.partition.csv import partition_csv
from unstructured.partition.xlsx import partition_xlsx
from unstructured.partition.docx import partition_docx
from unstructured.partition.html import partition_html
from unstructured.partition.text import partition_text
from unstructured.partition.md import partition_md
from unstructured.partition.pptx import partition_pptx
from unstructured.partition.ppt import partition_ppt
from unstructured.documents.elements import Text, Element

from app.utils.chunkingv2 import split_into_chunks
from app.utils.text_processing import elements_to_clean_text
from app.config import settings
from app.services.crm_chroma_service_v2 import get_crm_chroma_service_v2
from app.services.credit_service import get_credit_service

if TYPE_CHECKING:
    from app.services.storage_service import StorageService

logger = logging.getLogger(__name__)


# ============================================
# CONSTANTS
# ============================================

ALLOWED_KNOWLEDGE_EXTENSIONS = {"pdf", "docx", "txt", "csv"}
MAX_FILE_SIZE_MB = 10
MIN_TEXT_LENGTH = 50
MIN_WORD_COUNT = 20
MIN_ALPHA_RATIO = 0.3

# Extensions the processor can actually extract text from and embed.
# Everything else is stored in file manager (like Google Drive) but skipped for embedding.
EMBEDDABLE_EXTENSIONS = {
    # Documents
    "pdf", "docx", "doc", "txt", "csv", "md",
    "xlsx", "xls", "pptx", "ppt", "html", "htm", "log",
    # Images (via OCR proxy)
    "jpg", "jpeg", "png", "webp",
    # Audio / Video (via transcription proxy)
    "mp3", "wav", "m4a", "ogg", "flac",
    "mp4", "avi", "mov", "mkv", "webm",
}


class EmbeddingNotSupportedError(ValueError):
    """Raised when a file type is valid for storage but cannot be embedded."""
    pass


# ============================================
# QUALITY METRICS
# ============================================

class DocumentQualityMetrics:
    """Tracks extraction quality for logging, validation, and metadata."""
    def __init__(self):
        self.char_count: int = 0
        self.word_count: int = 0
        self.token_count: int = 0
        self.content_hash: str = ""
        self.has_tables: bool = False
        self.has_images: bool = False
        self.extraction_method: str = ""
        self.page_count: int = 0

    def to_dict(self) -> Dict[str, Any]:
        return {
            "char_count": self.char_count,
            "word_count": self.word_count,
            "token_count": self.token_count,
            "content_hash": self.content_hash,
            "has_tables": self.has_tables,
            "has_images": self.has_images,
            "extraction_method": self.extraction_method,
            "page_count": self.page_count,
        }


# ============================================
# MAIN CLASS
# ============================================

class DocumentProcessorV2:
    """Service for processing documents with enhanced text quality for RAG"""

    def __init__(self, storage_service: Optional['StorageService'] = None):
        self.storage_service = storage_service
        self.logger = logging.getLogger(__name__)
        self.logger.info("📦 Using Document Processor V2 (Improved)")

        self.chroma_service = get_crm_chroma_service_v2()
        self.credit_service = get_credit_service()

        base = getattr(settings, "PROXY_BASE_URL", "http://localhost:6657")
        self.proxy_base_url = base.rstrip("/") if base else "http://localhost:6657"

        # Tokenizer for accurate token counting
        try:
            self.tokenizer = tiktoken.encoding_for_model(settings.OPENAI_MODEL)
        except Exception:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")

        self.logger.info(f"🔗 V2 Proxy Target: {self.proxy_base_url}")

    # ============================================
    # VALIDATION (called from endpoint before processing)
    # ============================================

    @staticmethod
    def validate_knowledge_file(filename: str, content: bytes) -> str:
        # 1. Check for empty files
        if not content or len(content) == 0:
            raise ValueError("File kosong atau rusak. Mendukung PDF, DOCX, TXT, CSV (Max 10MB per file).")

        # 2. Check extension
        if not filename or "." not in filename:
            raise ValueError("Mendukung PDF, DOCX, TXT, CSV (Max 10MB per file).")

        ext = filename.rsplit(".", 1)[-1].lower()
        if ext not in ALLOWED_KNOWLEDGE_EXTENSIONS:
            raise ValueError("Format tidak didukung. Mendukung PDF, DOCX, TXT, CSV (Max 10MB per file).")

        # 3. Check file size (Max 10MB)
        size_mb = len(content) / (1024 * 1024)
        if size_mb > MAX_FILE_SIZE_MB:
            raise ValueError("Ukuran file terlalu besar. Mendukung PDF, DOCX, TXT, CSV (Max 10MB per file).")

        return ext

    def validate_quality(self, text: str, metrics: DocumentQualityMetrics, filename: str):
        issues = []

        if metrics.char_count < MIN_TEXT_LENGTH:
            issues.append(f"Too short ({metrics.char_count} chars, min {MIN_TEXT_LENGTH})")

        if metrics.word_count < MIN_WORD_COUNT:
            issues.append(f"Too few words ({metrics.word_count}, min {MIN_WORD_COUNT})")

        if metrics.char_count > 0:
            alpha_ratio = sum(c.isalnum() for c in text) / metrics.char_count
            if alpha_ratio < MIN_ALPHA_RATIO:
                issues.append(f"Low content quality ({alpha_ratio:.0%} alphanumeric, min {MIN_ALPHA_RATIO:.0%})")

        if not text.strip():
            issues.append("Empty after normalization")

        if issues:
            msg = f"Document '{filename}' failed quality check: {'; '.join(issues)}"
            self.logger.warning(f"⚠️ {msg}")
            raise ValueError(msg)

    # Magic bytes for binary formats.
    # Text-based formats (csv, txt, html, htm, md, log) are excluded —
    # they have no fixed signature and can start with any character.
    _MAGIC_BYTES: dict = {
        # Documents
        "pdf":  [b"%PDF"],
        "docx": [b"PK"],
        "xlsx": [b"PK"],
        "pptx": [b"PK"],
        "xls":  [b"\xd0\xcf"],
        "ppt":  [b"\xd0\xcf"],
        "doc":  [b"\xd0\xcf"],
        # Images
        "png":  [b"\x89PNG"],
        "jpg":  [b"\xff\xd8\xff"],
        "jpeg": [b"\xff\xd8\xff"],
        "webp": [b"RIFF"],
        "gif":  [b"GIF8"],
        "bmp":  [b"BM"],
        "tiff": [b"II*\x00", b"MM\x00*"],
        # Audio / Video
        "mp3":  [b"ID3", b"\xff\xfb", b"\xff\xf3", b"\xff\xf2"],
        "mp4":  [b"\x00\x00\x00\x18ftyp", b"\x00\x00\x00\x1cftyp"],
        "wav":  [b"RIFF"],
        # Archives
        "zip":  [b"PK"],
        "gz":   [b"\x1f\x8b"],
        "7z":   [b"7z\xbc\xaf"],
        "rar":  [b"Rar!"],
        # Executables / binaries — validated so a renamed HTML as .exe is caught,
        # but these will still fail at the embedding step (no parser for them).
        "exe":  [b"MZ"],
        "dll":  [b"MZ"],
        # Disk images
        "iso":  [b"\x00\x00\x01\xba", b"\x00\x00\x01\xb3"],  # MPEG / ISO start markers
    }

    @staticmethod
    def validate_file_magic(content: bytes, filename: str) -> None:
        """
        Reject files whose content doesn't match their declared extension.
        Raises ValueError with a clear message so the worker marks it as
        failed without deleting the file from storage.
        """
        if not filename or "." not in filename:
            return
        ext = filename.rsplit(".", 1)[-1].lower()
        expected = DocumentProcessorV2._MAGIC_BYTES.get(ext)
        if expected is None:
            return  # text-based or unknown — skip check
        if not any(content[:len(sig)] == sig for sig in expected):
            raise ValueError(
                f"File '{filename}' content does not match its .{ext} extension. "
                f"It may be a renamed file. "
                f"Please re-export as a proper {ext.upper()} file and re-upload."
            )

    def generate_content_hash(self, text: str) -> str:
        normalized = re.sub(r'\s+', ' ', text.lower().strip())
        return hashlib.sha256(normalized.encode()).hexdigest()[:16]

    async def check_duplicate(self, content_hash: str, agent_id: str, supabase) -> Optional[str]:
        try:
            response = supabase.table("knowledge_documents") \
                .select("name, metadata") \
                .eq("agent_id", agent_id) \
                .execute()

            if response.data:
                for doc in response.data:
                    existing_hash = (doc.get("metadata") or {}).get("content_hash")
                    if existing_hash and existing_hash == content_hash:
                        return doc.get("name", "unknown")
            return None
        except Exception as e:
            self.logger.warning(f"⚠️ Duplicate check failed (non-fatal): {e}")
            return None

    # ============================================
    # NORMALIZATION
    # ============================================

    def _normalize_text(self, text: str) -> str:
        if not text:
            return ""

        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r' {2,}', ' ', text)
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        text = text.replace('\u200b', '')
        text = text.replace('\ufeff', '')
        text = text.replace('\xa0', ' ')
        text = re.sub(r'\n\s*\d{1,3}\s*\n', '\n', text)

        return text.strip()

    # ============================================
    # MAIN ENTRY: process_document
    # ============================================

    def process_document(
        self,
        content: bytes,
        filename: str,
        folder_path: str,
        organization_id: str,
        file_id: str
    ) -> Tuple[str, DocumentQualityMetrics]:

        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        metrics = DocumentQualityMetrics()

        # Reject files whose bytes don't match their declared extension
        self.validate_file_magic(content, filename)

        # File types not in EMBEDDABLE_EXTENSIONS are stored (Google Drive style)
        # but skipped for embedding — raise so the worker marks them gracefully.
        if ext and ext not in EMBEDDABLE_EXTENSIONS:
            raise EmbeddingNotSupportedError(
                f"'.{ext}' files are stored but not supported for embedding. "
                f"The file is safely kept in storage."
            )

        if ext in ("mp3", "wav", "m4a", "ogg", "flac", "jpg", "jpeg", "png", "webp","mp4", "avi", "mov", "mkv", "webm"):
            if ext in ("jpg","jpeg","png","webp"):
                mime = mimetypes.guess_type(filename)[0] or "image/png"
                b64 = base64.b64encode(content).decode("ascii")
                image_data_url = f"data:{mime};base64,{b64}"
                headers = {"Content-Type": "application/json"}
                data = {"image_url": image_data_url}
                api_url = f"{self.proxy_base_url}/image/ocr"

                resp = requests.post(api_url, headers=headers, json=data, timeout=60)
                resp.raise_for_status()

                ocr_text = resp.json().get("content", "") or resp.json().get("text", "")

                # **CRUCIAL: Hitung metrics seperti PDF**
                normalized_text = self._normalize_text(ocr_text)
                metrics.char_count = len(normalized_text)
                metrics.word_count = len(normalized_text.split())
                try:
                    metrics.token_count = len(self.tokenizer.encode(normalized_text))
                except Exception:
                    metrics.token_count = metrics.word_count
                metrics.content_hash = self.generate_content_hash(normalized_text)
                metrics.extraction_method = "ocr_vision"
                
                self.logger.info(
                    f"✅ OCR '{filename}': {metrics.word_count} words, "
                    f"{metrics.token_count} tokens"
                )
                return normalized_text, metrics 

            if ext in ("mp3", "wav", "m4a", "ogg", "flac"):
                mime = mimetypes.guess_type(filename)[0] or "audio/mpeg"
                b64 = base64.b64encode(content).decode("ascii")
                audio_data_url = f"data:{mime};base64,{b64}"
                
                headers = {"Content-Type": "application/json"}
                data = {"url": audio_data_url}
                api_url = f"{self.proxy_base_url}/audio"
                
                resp = requests.post(api_url, headers=headers, json=data, timeout=300)
                resp.raise_for_status()
                
                transcription = resp.json()["output"]["result"]
                normalized_text = self._normalize_text(transcription)
                
                # Sama metrics calculation
                metrics.char_count = len(normalized_text)
                metrics.word_count = len(normalized_text.split())
                try:
                    metrics.token_count = len(self.tokenizer.encode(normalized_text))
                except:
                    metrics.token_count = metrics.word_count
                metrics.content_hash = self.generate_content_hash(normalized_text)
                metrics.extraction_method = "whisper_transcription"
                
                return normalized_text, metrics
            
            if ext in ("mp4", "avi", "mov", "mkv", "webm"):
                try:
                    # 1. Buat temp files
                    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as video_file:
                        video_path = video_file.name
                        video_file.write(content)
                    
                    audio_path = video_path.rsplit(".", 1)[0] + ".mp3"
                    
                    # 2. FFmpeg convert VIDEO → MP3 (silent, fast)
                    cmd = [
                        "ffmpeg", "-y", "-i", video_path, 
                        "-vn", "-acodec", "mp3", "-ar", "16000", 
                        audio_path, "-loglevel", "quiet"
                    ]
                    result = subprocess.run(cmd, capture_output=True, timeout=120)
                    result.check_returncode()
                    
                    # 3. Baca audio bytes → kirim ke /audio
                    with open(audio_path, "rb") as audio_file:
                        audio_bytes = audio_file.read()
                    
                    mime = "audio/mpeg"
                    b64 = base64.b64encode(audio_bytes).decode("ascii")
                    audio_data_url = f"data:{mime};base64,{b64}"
                    
                    # 4. Proxy /audio (pakai route yang SUDAH ADA!)
                    headers = {"Content-Type": "application/json"}
                    data = {"url": audio_data_url}
                    api_url = f"{self.proxy_base_url}/audio"
                    
                    resp = requests.post(api_url, headers=headers, json=data, timeout=300)
                    resp.raise_for_status()
                    
                    transcription = resp.json()["output"]["result"]
                    
                    # 5. Metrics
                    normalized_text = self._normalize_text(transcription)
                    metrics.char_count = len(normalized_text)
                    metrics.word_count = len(normalized_text.split())
                    try:
                        metrics.token_count = len(self.tokenizer.encode(normalized_text))
                    except:
                        metrics.token_count = metrics.word_count
                    metrics.content_hash = self.generate_content_hash(normalized_text)
                    metrics.extraction_method = "video_to_audio_whisper"
                    
                    self.logger.info(
                        f"✅ Video '{filename}' → Audio → Whisper: {metrics.word_count} words"
                    )
                    
                    # 6. Cleanup
                    os.unlink(video_path)
                    os.unlink(audio_path)
                    
                    return normalized_text, metrics
                except subprocess.CalledProcessError as e:
                    self.logger.error(f"❌ FFmpeg failed for {filename}: {e}")
                    return "[Video conversion failed]", metrics
                except Exception as e:
                    self.logger.error(f"❌ Video processing error: {e}")
                    if 'video_path' in locals():
                        os.unlink(video_path)
                    if 'audio_path' in locals() and os.path.exists(audio_path):
                        os.unlink(audio_path)
                    return f"[Error: {str(e)}]", metrics
            return "", metrics

        if ext in ("pdf", "docx", "csv", "txt"):
            raw_text, metrics = self._extract_lightweight(content, ext)
            normalized_text = self._normalize_text(raw_text)

            metrics.char_count = len(normalized_text)
            metrics.word_count = len(normalized_text.split())
            try:
                metrics.token_count = len(self.tokenizer.encode(normalized_text))
            except Exception:
                metrics.token_count = metrics.word_count
            metrics.content_hash = self.generate_content_hash(normalized_text)

            # self.logger.info(
            #     f"✅ Extracted '{filename}': {metrics.word_count} words, "
            #     f"{metrics.token_count} tokens, method={metrics.extraction_method}"
            # )

            return normalized_text, metrics

        elements = self._partition_by_type(content, filename, ext)
        clean_text = elements_to_clean_text(elements)
        normalized_text = self._normalize_text(clean_text)

        metrics.char_count = len(normalized_text)
        metrics.word_count = len(normalized_text.split())
        try:
            metrics.token_count = len(self.tokenizer.encode(normalized_text))
        except Exception:
            metrics.token_count = metrics.word_count
        metrics.content_hash = self.generate_content_hash(normalized_text)
        metrics.extraction_method = "unstructured_fallback"

        if any(el.category == "Table" for el in elements):
            metrics.has_tables = True

        return normalized_text, metrics

    # ============================================
    # LIGHTWEIGHT PARSERS
    # ============================================

    def _extract_lightweight(self, content: bytes, ext: str) -> Tuple[str, DocumentQualityMetrics]:
        if ext == "pdf":
            return self._extract_pdf(content)
        elif ext == "docx":
            return self._extract_docx(content)
        elif ext == "csv":
            return self._extract_csv(content)
        elif ext == "txt":
            return self._extract_txt(content)
        else:
            raise ValueError(f"No lightweight parser for '.{ext}'")

    # ============================================
    # PDF EXTRACTION (hi_res with pdfplumber fallback)
    # ============================================

    def _extract_pdf(self, content: bytes) -> Tuple[str, DocumentQualityMetrics]:
        metrics = DocumentQualityMetrics()

        try:
            self.logger.info("📄 PDF: Using hi_res strategy (layout + tables + images)")

            timeout_secs = getattr(settings, "PDF_EXTRACTION_TIMEOUT", 60)

            # Run partition_pdf in a thread so we can enforce a wall-clock timeout.
            # A corrupted or adversarial PDF can otherwise stall the worker indefinitely.
            def _run_partition(tmpdir: str):
                return partition_pdf(
                    file=io.BytesIO(content),
                    strategy="hi_res",
                    infer_table_structure=False,
                    extract_images_in_pdf=True,
                    extract_image_block_types=["Image"],
                    extract_image_block_output_dir=tmpdir,
                    languages=["ind", "eng"],
                    max_partition=None,
                )

            # 1. OPEN THE SANDBOX
            with tempfile.TemporaryDirectory() as tmpdir:
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(_run_partition, tmpdir)
                    try:
                        elements = future.result(timeout=timeout_secs)
                    except concurrent.futures.TimeoutError:
                        self.logger.warning(
                            f"⚠️ PDF hi_res extraction timed out after {timeout_secs}s, "
                            f"falling back to pdfplumber"
                        )
                        return self._extract_pdf_pdfplumber(content)

                if not elements:
                    return self._extract_pdf_pdfplumber(content)

                text_blocks = []
                element_counts = {"text": 0, "table": 0, "image": 0, "title": 0, "other": 0}

                # 2. THIS LOOP MUST BE INDENTED INSIDE THE 'WITH' BLOCK
                for el in elements:
                    category = getattr(el, "category", "").lower() if hasattr(el, "category") else ""
                    el_text = getattr(el, "text", "") or ""
                    el_metadata = getattr(el, "metadata", None)

                    if category in ("title", "header"):
                        element_counts["title"] += 1
                        if el_text.strip():
                            text_blocks.append(f"\n## {el_text.strip()}\n")

                    elif category == "table":
                        element_counts["table"] += 1
                        metrics.has_tables = True
                        if el_text.strip():
                            text_blocks.append(f"\n[Table]\n{el_text.strip()}\n")

                    elif category == "image":
                        element_counts["image"] += 1
                        metrics.has_images = True
                        
                        # 3. GET THE IMAGE PATH
                        image_path = getattr(el_metadata, "image_path", None) if el_metadata else None
                        
                        # 4. READ IT, ENCODE IT, AND SEND TO OCR PROXY
                        if image_path and os.path.exists(image_path):
                            with open(image_path, "rb") as img_file:
                                img_bytes = img_file.read()
                            
                            mime, _ = mimetypes.guess_type(image_path)
                            mime = mime or "image/png"
                            b64 = base64.b64encode(img_bytes).decode("ascii")
                            image_data_url = f"data:{mime};base64,{b64}"

                            try:
                                api_url = f"{self.proxy_base_url}/image/ocr"
                                headers = {"Content-Type": "application/json"}
                                data = {"image_url": image_data_url}

                                resp = requests.post(api_url, headers=headers, json=data, timeout=60)
                                resp.raise_for_status()

                                ocr_text = resp.json().get("content", "") or resp.json().get("text", "")
                                if ocr_text.strip():
                                    text_blocks.append(f"\n[Extracted Image Data]\n{ocr_text.strip()}\n")
                            except Exception as ocr_err:
                                self.logger.warning(f"⚠️ OCR proxy failed: {ocr_err}")
                        
                        # 5. FALLBACK IF NO IMAGE WAS SAVED
                        elif el_text.strip() and len(el_text.strip()) > 5:
                            text_blocks.append(f"\n[Figure: {el_text.strip()}]\n")

                    else:
                        if el_text.strip():
                            element_counts["text"] += 1
                            text_blocks.append(el_text.strip())

                # Count pages
                page_numbers = set()
                for el in elements:
                    el_metadata = getattr(el, "metadata", None)
                    if el_metadata:
                        pn = getattr(el_metadata, "page_number", None)
                        if pn:
                            page_numbers.add(pn)

                metrics.page_count = len(page_numbers) if page_numbers else 0
                metrics.extraction_method = "unstructured_hi_res_with_ocr"

                result = "\n\n".join(text_blocks)

            # --- THE 'WITH' BLOCK ENDS HERE. THE TEMP DIR IS NOW AUTOMATICALLY DELETED. ---

            if len(result.strip()) < MIN_TEXT_LENGTH:
                self.logger.warning(
                    f"⚠️ hi_res extracted only {len(result.strip())} chars, "
                    f"falling back to pdfplumber"
                )
                return self._extract_pdf_pdfplumber(content)

            return result, metrics

        except Exception as e:
            self.logger.error(f"⚠️ hi_res extraction failed: {e}")
            self.logger.info("🔄 Falling back to pdfplumber...")
            return self._extract_pdf_pdfplumber(content)
        
    def _extract_pdf_pdfplumber(self, content: bytes) -> Tuple[str, DocumentQualityMetrics]:
        """
        Improved fallback PDF extraction using pdfplumber.
        """
        metrics = DocumentQualityMetrics()
        metrics.extraction_method = "pdfplumber_improved"
        text_blocks = []
        current_section = ""        # Track section headers across pages
        prev_table_headers = []     # For cross-page table continuation
        extracted_images = []       # Collect base64 images

        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                metrics.page_count = len(pdf.pages)

                for page_idx, page in enumerate(pdf.pages):
                    page_parts = []

                    # ──────────────────────────────────
                    # STEP 1: Extract & format tables
                    # ──────────────────────────────────
                    table_bboxes = []
                    table_finder = page.find_tables()
                    if table_finder:
                        table_bboxes = [t.bbox for t in table_finder]

                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if not table:
                                continue

                            # --- Skip single-column "wrapper" tables ---
                            # pdfplumber sometimes wraps the entire page in a
                            # 1-column table. These are layout artifacts.
                            max_cols = max(len(row) for row in table if row)
                            if max_cols <= 1:
                                continue

                            # --- Filter blob rows ---
                            clean_rows = []
                            for row in table:
                                if not row:
                                    continue
                                non_none = [c for c in row if c and c.strip()]
                                is_blob = (
                                    len(non_none) == 1
                                    and max_cols > 1
                                    and len(non_none[0]) > 200
                                )
                                if is_blob:
                                    continue
                                clean_rows.append(row)

                            if not clean_rows:
                                continue

                            metrics.has_tables = True

                            # --- Detect headers ---
                            headers = []
                            data_start = 0

                            first_row = clean_rows[0]
                            first_cells = [str(c).strip().replace('\n', ' ') if c else "" for c in first_row]

                            # Check if this is a continuation table:
                            # first cell is empty → data continues from previous page
                            is_continuation = (
                                not first_cells[0]
                                and any(c for c in first_cells[1:])
                                and prev_table_headers
                                and len(prev_table_headers) == len(first_cells)
                            )

                            if is_continuation:
                                headers = prev_table_headers
                                data_start = 0  # All rows are data
                            else:
                                # Check if the first row looks like a header
                                # (short text, no sentences, distinct from data rows)
                                avg_header_len = sum(len(c) for c in first_cells if c) / max(len([c for c in first_cells if c]), 1)
                                if avg_header_len < 60:  # Header cells are typically short
                                    headers = first_cells
                                    data_start = 1
                                else:
                                    data_start = 0

                            # Save headers for potential cross-page continuation
                            if headers:
                                prev_table_headers = headers

                            # --- Build formatted output ---
                            table_lines = ["\n[Table]"]

                            visible_headers = [h for h in headers if h]
                            if visible_headers:
                                table_lines.append(" | ".join(visible_headers))

                            for row in clean_rows[data_start:]:
                                cells = [
                                    str(c).strip().replace('\n', ' ') if c else ""
                                    for c in row
                                ]

                                # Deduplicate merged cells
                                deduped = []
                                for i, c in enumerate(cells):
                                    if i == 0 or c != cells[i - 1]:
                                        deduped.append(c)
                                    else:
                                        deduped.append("")
                                cells = deduped

                                # Key-value format when headers available
                                if headers and len(headers) == len(cells):
                                    parts = []
                                    for h, c in zip(headers, cells):
                                        if c and h:
                                            parts.append(f"{h}: {c}")
                                        elif c:
                                            parts.append(c)
                                    row_text = ". ".join(parts)
                                else:
                                    row_text = " | ".join(c for c in cells if c)

                                if row_text.strip():
                                    table_lines.append(row_text)

                            if len(table_lines) > 1:
                                page_parts.extend(table_lines)

                    # ──────────────────────────────────
                    # STEP 2: Extract text EXCLUDING table regions
                    # ──────────────────────────────────
                    if table_bboxes:
                        def _not_in_table(obj):
                            if obj.get("object_type") != "char":
                                return True
                            ox = obj.get("x0", 0)
                            ot = obj.get("top", 0)
                            for (x0, top, x1, bottom) in table_bboxes:
                                if x0 - 2 <= ox <= x1 + 2 and top - 2 <= ot <= bottom + 2:
                                    return False
                            return True

                        filtered_page = page.filter(_not_in_table)
                        page_text = filtered_page.extract_text(layout=False)
                    else:
                        page_text = page.extract_text(layout=False)

                    if page_text and page_text.strip():
                        # Track section headers
                        for line in page_text.strip().split('\n'):
                            stripped = line.strip()
                            if re.match(r'^\d+\.\s+[A-Z]', stripped) and len(stripped) < 80:
                                current_section = stripped

                        page_parts.append(page_text.strip())

                    # ──────────────────────────────────
                    # STEP 3: Extract images as base64
                    # ──────────────────────────────────
                    page_images = page.images
                    if page_images:
                        metrics.has_images = True

                        for img_idx, img_meta in enumerate(page_images):
                            try:
                                x0 = img_meta.get("x0", 0)
                                top_coord = img_meta.get("top", 0)
                                x1 = img_meta.get("x1", 0)
                                bottom_coord = img_meta.get("bottom", 0)

                                width = x1 - x0
                                height = bottom_coord - top_coord

                                # Skip tiny images (icons, bullets, decorations)
                                if width < 50 or height < 30:
                                    continue

                                # Crop with small padding and render at 200 DPI
                                bbox = (
                                    max(0, x0 - 5),
                                    max(0, top_coord - 5),
                                    min(page.width, x1 + 5),
                                    min(page.height, bottom_coord + 5),
                                )
                                cropped = page.crop(bbox)
                                rendered = cropped.to_image(resolution=200)

                                buf = io.BytesIO()
                                rendered.save(buf, format="PNG")
                                buf.seek(0)
                                img_bytes = buf.read()
                                b64_str = base64.b64encode(img_bytes).decode("ascii")

                                extracted_images.append({
                                    "page": page_idx + 1,
                                    "index": img_idx,
                                    "width": int(width),
                                    "height": int(height),
                                    "mime_type": "image/png",
                                    "base64": b64_str,
                                })

                                page_parts.append(
                                    f"\n[Image on page {page_idx + 1}: "
                                    f"diagram/figure {img_idx + 1}, "
                                    f"{int(width)}x{int(height)}px — "
                                    f"base64 image stored in metadata]\n"
                                )

                                self.logger.info(
                                    f"📸 Extracted image p{page_idx + 1}: "
                                    f"{int(width)}x{int(height)}px, "
                                    f"{len(b64_str)} b64 chars"
                                )

                            except Exception as img_err:
                                self.logger.warning(
                                    f"⚠️ Image extract failed p{page_idx + 1}: {img_err}"
                                )

                    # ──────────────────────────────────
                    # STEP 4: Cross-page section context
                    # ──────────────────────────────────
                    if page_parts and current_section and page_idx > 0:
                        first_line = page_parts[0].strip().split('\n')[0] if page_parts else ""
                        if not re.match(r'^\d+\.\s+[A-Z]', first_line):
                            page_parts.insert(0, f"[Continued: {current_section}]")

                    if page_parts:
                        text_blocks.append("\n".join(page_parts))

            # Store images in metrics for downstream use (e.g., OCR, embedding)
            if extracted_images:
                metrics.extracted_images = extracted_images  # type: ignore[attr-defined]

            result = "\n\n".join(text_blocks)

            return result, metrics

        except Exception as e:
            self.logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")

    # ============================================
    # HTML TABLE CONVERTER
    # ============================================

    def _html_table_to_text(self, html: str) -> str:
        """Convert HTML table to column-aware text for RAG."""
        try:
            soup = BeautifulSoup(html, "html.parser")
            table = soup.find("table")
            if not table:
                rows = soup.find_all("tr")
            else:
                rows = table.find_all("tr")

            if not rows:
                return soup.get_text(separator=" ", strip=True)

            header_cells = rows[0].find_all(["th", "td"])
            headers = [cell.get_text(strip=True) for cell in header_cells]

            lines = []

            if headers and any(h for h in headers):
                lines.append(" | ".join(headers))

            for row in rows[1:]:
                cells = [td.get_text(strip=True) for td in row.find_all(["td", "th"])]
                if not any(c for c in cells):
                    continue

                if headers and len(headers) == len(cells):
                    line = ", ".join(
                        f"{h}: {c}" for h, c in zip(headers, cells) if c
                    )
                else:
                    line = " | ".join(cells)

                if line.strip():
                    lines.append(line)

            return "\n".join(lines)

        except Exception as e:
            self.logger.warning(f"⚠️ HTML table parse failed: {e}")
            try:
                return BeautifulSoup(html, "html.parser").get_text(separator=" ", strip=True)
            except Exception:
                return html

    # ============================================
    # DOCX EXTRACTION (Enhanced with image OCR)
    # ============================================

    def _extract_docx(self, content: bytes) -> Tuple[str, DocumentQualityMetrics]:
        """
        DOCX extraction using python-docx + unstructured for images.
        1. python-docx: text with heading hierarchy + tables
        2. If images found: unstructured partition_docx for OCR
        """
        metrics = DocumentQualityMetrics()
        metrics.extraction_method = "python_docx"
        text_blocks = []

        try:
            doc = DocxDocument(io.BytesIO(content))

            # --- Check if document has images ---
            has_images = False
            try:
                for rel in doc.part.rels.values():
                    if "image" in str(getattr(rel, 'reltype', '')).lower():
                        has_images = True
                        break
            except Exception:
                pass

            # --- Extract paragraphs with heading hierarchy ---
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    prefix = "#" * min(level, 4)
                    text_blocks.append(f"\n{prefix} {text}\n")
                elif style_name == "List Paragraph" or style_name.startswith("List"):
                    text_blocks.append(f"• {text}")
                else:
                    text_blocks.append(text)

            # --- Extract tables with column context ---
            if doc.tables:
                metrics.has_tables = True
                for table in doc.tables:
                    rows = table.rows
                    if not rows:
                        continue

                    headers = [cell.text.strip() for cell in rows[0].cells]

                    # Deduplicate headers (merged cells produce duplicates)
                    seen_headers = []
                    for h in headers:
                        if h not in seen_headers:
                            seen_headers.append(h)
                        else:
                            seen_headers.append("")
                    headers = seen_headers

                    table_lines = ["\n[Table]"]
                    table_lines.append(" | ".join(h for h in headers if h))

                    for row in rows[1:]:
                        cells = [cell.text.strip() for cell in row.cells]

                        clean_cells = []
                        for i, c in enumerate(cells):
                            if i == 0 or c != cells[i - 1]:
                                clean_cells.append(c)
                            else:
                                clean_cells.append("")
                        cells = clean_cells

                        if headers and len(headers) == len(cells):
                            row_text = ", ".join(
                                f"{h}: {c}" for h, c in zip(headers, cells) if c and h
                            )
                        else:
                            row_text = " | ".join(c for c in cells if c)

                        if row_text.strip():
                            table_lines.append(row_text)

                    text_blocks.extend(table_lines)

            # --- Extract image text via unstructured ---
            if has_images:
                metrics.has_images = True
                try:
                    self.logger.info("🖼️ DOCX contains images, running OCR via unstructured...")
                    elements = partition_docx(
                        file=io.BytesIO(content),
                        infer_table_structure=False,
                    )

                    for el in elements:
                        category = getattr(el, "category", "").lower() if hasattr(el, "category") else ""
                        el_text = getattr(el, "text", "") or ""

                        if category == "image" and el_text.strip() and len(el_text.strip()) > 5:
                            text_blocks.append(f"\n[Figure: {el_text.strip()}]\n")

                    self.logger.info("✅ DOCX image OCR complete")
                    metrics.extraction_method = "python_docx+unstructured_ocr"

                except Exception as img_err:
                    self.logger.warning(f"⚠️ DOCX image OCR failed (non-fatal): {img_err}")

            return "\n".join(text_blocks), metrics

        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    # CSV EXTRACTION (unchanged)
    def _extract_csv(self, content: bytes) -> Tuple[str, DocumentQualityMetrics]:
        metrics = DocumentQualityMetrics()
        metrics.has_tables = True
        metrics.extraction_method = "pandas_csv"

        try:
            # ★ FIX: charset_normalizer returns ResultObject, not dict
            detected = detect_encoding(content)
            if isinstance(detected, dict):
                encoding = detected.get("encoding", "utf-8") or "utf-8"
            elif hasattr(detected, "best"):
                best = detected.best()
                encoding = best.encoding if best else "utf-8"
            else:
                encoding = "utf-8"
            

            df = pd.read_csv(io.BytesIO(content), encoding=encoding)
            if df.empty:
                return "", metrics

            columns = list(df.columns)
            
            # ★ Build a rich schema header that RAG can actually find
            schema_lines = [
                f"## CSV Data Overview",
                f"Total rows: {len(df)}. Total columns: {len(columns)}.",
                f"Column names: {', '.join(columns)}.",
            ]
            
            # ★ Add column types + sample values so LLM understands the data
            for col in columns:
                non_null = df[col].dropna()
                if non_null.empty:
                    schema_lines.append(f"- {col}: all empty")
                    continue
                dtype = str(df[col].dtype)
                sample = str(non_null.iloc[0])[:80]
                unique_count = non_null.nunique()
                schema_lines.append(f"- {col} ({dtype}): {unique_count} unique values, example: {sample}")
            
            text_blocks = ["\n".join(schema_lines)]

            # ★ FIX: Section markers now cover ALL rows including first batch
            ROWS_PER_SECTION = 50
            total_rows = len(df)
            
            for idx, (_, row) in enumerate(df.iterrows()):
                # Insert section header at the START of each batch
                if idx % ROWS_PER_SECTION == 0:
                    section_end = min(idx + ROWS_PER_SECTION, total_rows)
                    text_blocks.append(f"\n## Rows {idx + 1}-{section_end}\n")
                
                parts = []
                for col in columns:
                    val = row[col]
                    if pd.notna(val):
                        val_str = str(val).strip()
                        if val_str:
                            parts.append(f"{col}: {val_str}")
                if parts:
                    text_blocks.append(". ".join(parts))

            return "\n".join(text_blocks), metrics

        except Exception as e:
            self.logger.error(f"CSV extraction failed: {e}")
            raise ValueError(f"Failed to extract text from CSV: {e}")
    
    # ============================================
    # TXT EXTRACTION (unchanged)
    # ============================================

    def _extract_txt(self, content: bytes) -> Tuple[str, DocumentQualityMetrics]:
        metrics = DocumentQualityMetrics()
        metrics.extraction_method = "charset_txt"
        try:
            detected = detect_encoding(content)
            if isinstance(detected, dict):
                encoding = detected.get("encoding", "utf-8") or "utf-8"
            elif hasattr(detected, "best"):
                best = detected.best()
                encoding = best.encoding if best else "utf-8"
            else:
                encoding = "utf-8"
            
            return content.decode(encoding, errors="replace"), metrics
        except Exception:
            return content.decode("utf-8", errors="ignore"), metrics

    # ============================================
    # UNSTRUCTURED FALLBACK (for html, md, pptx, xlsx, etc.)
    # ============================================

    def _partition_by_type(self, content: bytes, filename: str, ext: str) -> List[Element]:
        fobj = io.BytesIO(content)

        try:
            if ext == "csv":
                return partition_csv(file=fobj)
            elif ext in ("xlsx", "xls"):
                # Magic bytes already validated in process_document before we get here.
                # .xlsx → PK (ZIP), .xls → D0CF (OLE2 legacy)
                return partition_xlsx(file=fobj)
            elif ext == "pdf":
                return partition_pdf(
                    file=fobj,
                    strategy="hi_res",
                    infer_table_structure=True,
                    extract_images_in_pdf=True,
                    languages=["ind", "eng"],
                )
            elif ext == "docx":
                return partition_docx(file=fobj)
            elif ext in ("html", "htm"):
                return partition_html(file=fobj)
            elif ext in ("md", "markdown"):
                return partition_md(file=fobj)
            elif ext in ("txt", "log"):
                return partition_text(file=fobj)
            elif ext == "pptx":
                return partition_pptx(file=fobj)
            elif ext == "ppt":
                return partition_ppt(file=fobj)
            else:
                with tempfile.NamedTemporaryFile(delete=True, suffix=f".{ext}") as tmp:
                    tmp.write(content)
                    tmp.flush()
                    return partition(filename=tmp.name, strategy="auto")

        except Exception as e:
            self.logger.error(f"❌ Local processing failed for {filename}: {e}")
            raise

    # ============================================
    # PROXY HANDLERS (image OCR + audio)
    # ============================================

    def _process_image(self, content: bytes, filename: str, folder_path: str, organization_id: str, file_id: str) -> Tuple[str, List[dict]]:
        try:
            self.logger.info(f"🌆 Processing image file via V2: {filename}")
            mime, _ = mimetypes.guess_type(filename)
            if not mime or not mime.startswith("image/"):
                mime = "image/png"

            b64 = base64.b64encode(content).decode("ascii")
            image_url = f"data:{mime};base64,{b64}"

            api_url = f"{self.proxy_base_url}/image/ocr"
            headers = {"Content-Type": "application/json"}
            data = {"image_url": image_url}

            resp = requests.post(api_url, headers=headers, json=data, timeout=60)
            resp.raise_for_status()

            result = resp.json()
            text = result.get("content", "") or result.get("text", "")
            text = self._normalize_text(text)

            element = Text(text=text)
            return text, self._elements_to_json([element])

        except Exception as e:
            self.logger.error(f"❌ Image OCR failed: {e}")
            return "", []

    def _process_audio(self, content: bytes, filename: str, folder_path: str, organization_id: str, file_id: str) -> Tuple[str, List[dict]]:
        try:
            self.logger.warning(f"⚠️ Audio processing not fully implemented for {filename}")
            return "", []
        except Exception as e:
            self.logger.error(f"❌ Audio transcription failed: {e}")
            return "", []

    # ============================================
    # LEGACY METHODS (kept for backward compat)
    # ============================================

    async def process_and_embed(self, agent_id: str, organization_id: str, file_content: bytes, file_type: str, filename: str):
        try:
            self.logger.info(f"⚙️ Processing file: {filename}")

            text = self._extract_text(file_content, filename)
            if not text:
                return {"success": False, "error": "No text extracted"}

            text = self._normalize_text(text)

            if len(text) < 50:
                return {"success": False, "error": "Document too short after normalization"}

            chunks = split_into_chunks(text, size=512, overlap=100)

            self.logger.info(f"✂️ Generated {len(chunks)} chunks (avg: {len(text) // max(len(chunks), 1)} chars/chunk)")

            metadatas = [{"source": filename, "doc_id": filename} for _ in chunks]

            result = await self.chroma_service.add_documents(
                agent_id=agent_id,
                texts=chunks,
                metadatas=metadatas
            )

            if result.get("success"):
                try:
                    usage = result.get("usage", {})
                    # Calculate raw tokens from the proxy response
                    total_tokens = int(usage.get("total_tokens", 0))
                    
                    if total_tokens > 0 and organization_id:
                        import math
                        from app.models.credit import CreditUsageCreate, QueryType, QueryStatus
                        from app.services.subscription_service import get_subscription_service
                        
                        # 1. APPLY THE EXCHANGE RATE: 1 Subscription Credit = 250 Tokens
                        credits_to_deduct = math.ceil(total_tokens / 250)
                        
                        # Internal cost tracking
                        cost = float(usage.get("cost_usd", total_tokens * 0.0000002))

                        self.logger.info(f"💸 Deducting {credits_to_deduct} credits for embedding {len(chunks)} chunks...")
                        
                        # 2. Construct the strict Pydantic payload for File Uploads
                        usage_payload = CreditUsageCreate(
                            organization_id=organization_id,
                            query_type=QueryType.UPLOAD_FILE,
                            query_text=f"Knowledge Embedding ({len(chunks)} chunks)",
                            credits_used=credits_to_deduct,
                            status=QueryStatus.COMPLETED,
                            input_tokens=total_tokens,
                            output_tokens=0,
                            cost=cost,
                            metadata={"agent_id": agent_id, "file": filename, "chunk_count": len(chunks)}
                        )

                        # 3. Write immutable record to the Ledger
                        await self.credit_service.log_usage(usage_payload)
                        
                        # 4. Enforce the limit against the Subscription
                        if credits_to_deduct > 0:
                            sub_service = get_subscription_service()
                            await sub_service.increment_usage(organization_id, credits_to_deduct)

                    else:
                        self.logger.info("🆓 Embedding cost was 0 tokens.")

                except Exception as bill_err:
                    self.logger.error(f"🚨 BILLING FAILURE for {organization_id}: {bill_err}")

                return {"success": True, "chunks": len(chunks)}

            else:
                return {"success": False, "error": result.get("error", "Chroma storage failed")}

        except Exception as e:
            self.logger.error(f"❌ Processing Failed: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    def _extract_text(self, content: bytes, filename: str) -> str:
        try:
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
            elements = self._partition_by_type(content, filename, ext)
            raw_text = elements_to_clean_text(elements)
            return self._normalize_text(raw_text)
        except Exception as e:
            self.logger.error(f"Text extraction failed for {filename}: {e}")
            return ""

    @staticmethod
    def _elements_to_json(elements: List[Element]) -> List[dict]:
        result = []
        for el in elements:
            element_dict = {
                "category": getattr(el, "category", None),
                "text": getattr(el, "text", None)
            }
            metadata = getattr(el, "metadata", None)
            element_dict["metadata"] = (
                metadata.to_dict() if hasattr(metadata, "to_dict")
                else (metadata if isinstance(metadata, dict) else None)
            )
            result.append(element_dict)
        return result